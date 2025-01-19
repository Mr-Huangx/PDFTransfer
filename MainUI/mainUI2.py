import sys
import os
import tempfile
import warnings
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QVBoxLayout, QScrollArea, QWidget, QLabel,
    QLineEdit, QPushButton, QHBoxLayout, QMenuBar, QFileDialog, QMessageBox
)
from PyQt5.QtCore import Qt, QTimer, QUrl  # 导入 QUrl
from PyQt5.QtGui import QFont
from PyQt5.QtWebEngineWidgets import QWebEngineView, QWebEngineSettings  # 用于显示 PDF
from docx import Document
from docx2pdf import convert  # 用于将 Word 转换为 PDF
import re
from PyQt5.QtCore import QLoggingCategory
QLoggingCategory.setFilterRules("js=false")


# 忽略 DeprecationWarning
warnings.filterwarnings("ignore", category=DeprecationWarning)


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        # 设置窗口标题和初始大小
        self.setWindowTitle("自动文档填充")
        self.setGeometry(100, 100, 1600, 900)

        # 初始化 UI
        self.init_ui()

        # 初始状态隐藏主页面
        self.central_widget.hide()

    def init_ui(self):
        # 添加菜单栏
        menu_bar = QMenuBar(self)
        file_menu = menu_bar.addMenu("文件")
        load_template_action = file_menu.addAction("加载模板")
        load_template_action.triggered.connect(self.load_template)
        self.setMenuBar(menu_bar)

        # 创建主布局
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        main_layout = QHBoxLayout(self.central_widget)

        # 左侧固定宽度区域（包含表单和提交按钮）
        self.left_widget = QWidget()
        self.left_widget.setFixedWidth(500)  # 固定宽度为 500
        self.left_layout = QVBoxLayout(self.left_widget)

        # 创建滚动视图
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.form_container = QWidget()
        self.form_layout = QVBoxLayout(self.form_container)

        # 添加表单内容到滚动视图
        self.inputs = {}
        self.scroll_area.setWidget(self.form_container)

        # 将滚动视图添加到左侧布局
        self.left_layout.addWidget(self.scroll_area)

        # 创建按钮布局（水平布局）
        button_layout = QHBoxLayout()

        # 添加“更新预览”按钮
        self.update_preview_button = QPushButton("更新预览")
        self.update_preview_button.setFont(QFont("Segoe UI", 10))
        self.update_preview_button.setStyleSheet(
            "QPushButton {"
            "padding: 10px;"
            "background-color: #28a745;"
            "color: white;"
            "border: none;"
            "border-radius: 4px;"
            "}"
            "QPushButton:hover {"
            "background-color: #218838;"
            "}"
        )
        self.update_preview_button.clicked.connect(self.update_preview)
        button_layout.addWidget(self.update_preview_button)

        # 添加“生成文档”按钮
        self.generate_button = QPushButton("生成文档")
        self.generate_button.setFont(QFont("Segoe UI", 10))
        self.generate_button.setStyleSheet(
            "QPushButton {"
            "padding: 10px;"
            "background-color: #007BFF;"
            "color: white;"
            "border: none;"
            "border-radius: 4px;"
            "}"
            "QPushButton:hover {"
            "background-color: #0056b3;"
            "}"
        )
        self.generate_button.clicked.connect(self.fill_word_template)
        button_layout.addWidget(self.generate_button)

        # 将按钮布局添加到左侧布局
        self.left_layout.addLayout(button_layout)

        # 右侧预览区域（PDF 预览）
        self.preview = QWebEngineView()  # 使用 QWebEngineView 显示 PDF
        # self.preview = DisabledWebEngineView()
        self.preview.setStyleSheet(
            "QWebEngineView {"
            "padding: 10px;"
            "border: 1px solid #ccc;"
            "border-radius: 4px;"
            "background-color: #f5f5f5;"
            "}"
        )

        # 启用 PDF 插件
        # self.preview.settings().setAttribute(QWebEngineSettings.PdfViewerEnabled, True)
        self.preview.settings().setAttribute(QWebEngineSettings.PluginsEnabled, True)

        
        # 将左侧和右侧添加到主布局
        main_layout.addWidget(self.left_widget)
        main_layout.addWidget(self.preview)

    def load_template(self):
        """加载 Word 模板文件"""
        file_path, _ = QFileDialog.getOpenFileName(self, "打开 Word 模板", "", "Word 文件 (*.docx)")
        if file_path:
            # 显示主页面
            self.central_widget.show()

            # 加载模板
            try:
                self.template_document = Document(file_path)
                print("模板文件加载成功。")
                self.template_document_path = file_path # 记录模板文件的地址

            except Exception as e:
                print("模板文件加载失败:", str(e))
                return

            # 动态生成表单字段
            self.generate_form_fields()

            # 生成 PDF 预览
            self.generate_preview_pdf()

    def generate_form_fields(self):
        """从模板中读取占位符并动态生成表单字段"""
        # 清空现有表单
        for i in reversed(range(self.form_layout.count())):
            self.form_layout.itemAt(i).widget().setParent(None)
        self.inputs.clear()

        # 读取模板中的占位符
        pattern = re.compile(r"\{\{.*?\}\}")  # 匹配占位符的正则表达式
        placeholders = list()
        for paragraph in self.template_document.paragraphs:
            matches = pattern.findall(paragraph.text)
            for match in matches:
                if match not in placeholders:
                    placeholders.append(match)

        # 从表格中提取占位符
        for table in self.template_document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = pattern.findall(paragraph.text)
                        for match in matches:
                            if match not in placeholders:  # 避免重复
                                placeholders.append(match)

        # 动态生成表单字段
        font_label = QFont("Segoe UI", 10)
        font_input = QFont("Segoe UI", 10)

        for placeholder in placeholders:
            label = QLabel(placeholder[2:-2])
            label.setFont(font_label)
            input_box = QLineEdit()
            input_box.setFont(font_input)
            input_box.setStyleSheet(
                "QLineEdit {"
                "padding: 8px;"
                "border: 1px solid #ccc;"
                "border-radius: 4px;"
                "background-color: #f9f9f9;"
                "}"
                "QLineEdit:focus {"
                "border: 1px solid #007BFF;"
                "background-color: #fff;"
                "}"
            )
            self.inputs[placeholder] = input_box
            self.form_layout.addWidget(label)
            self.form_layout.addWidget(input_box)

    def generate_preview_pdf(self):
        """生成 PDF 并在预览区域显示"""
        if not hasattr(self, 'template_document'):
            print("错误：未加载模板文件。")
            return

        # 创建临时 Word 文件
        # temp_docx_path = os.path.join(tempfile.gettempdir(), "temp_template.docx")
        # print("临时 Word 文件路径:", temp_docx_path)
        # self.template_document.save(temp_docx_path)

        # 将 Word 转换为 PDF
        self.temp_pdf_path = os.path.join(tempfile.gettempdir(), "temp_preview.pdf")
        print("临时 PDF 文件路径:", self.temp_pdf_path)
        try:
            convert(self.template_document_path, self.temp_pdf_path)
            print("Word 转换为 PDF 成功。")
        except Exception as e:
            print("Word 转换为 PDF 失败:", str(e))
            return

        # 检查 PDF 文件是否存在
        if not os.path.exists(self.temp_pdf_path):
            print("错误：PDF 文件未生成。")
            return

        # 在预览区域显示 PDF
        pdf_url = QUrl.fromLocalFile(self.temp_pdf_path)
        print("PDF 文件 URL:", pdf_url.toString())
        self.preview.setUrl(pdf_url)

        # 检查 QWebEngineView 是否加载成功
        self.preview.loadFinished.connect(self.on_load_finished)

    def on_load_finished(self, success):
        """PDF 加载完成时的回调函数"""
        if success:
            print("PDF 文件加载回调成功。")
        else:
            print("PDF 文件加载回调失败。")

    def update_preview(self):
        """更新预览"""
        # 第一步，打开document文件
        # 第二步，替换占位符
        # 第三步，生成临时文件word
        # 第四步，展示临时文件PDF
        doc = Document(self.template_document_path)

        # 替换段落中的占位符
        for paragraph in doc.paragraphs:
            for field, input_box in self.inputs.items():
                placeholder = field
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, input_box.text())

        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, input_box in self.inputs.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, input_box.text())

        # 先转换成临时word
        temp_docx_path = os.path.join(tempfile.gettempdir(), "temp_template.docx")
        print(f"更新预览生成的temp_docx_path:{temp_docx_path}")

        # 转换成临时PDF
        # 删除原始临时PDF
        os.remove(self.temp_pdf_path)
        
        convert(temp_docx_path, self.temp_pdf_path)
        print("Word 转换为 PDF 成功。")
        print(f"更新预览生成的temp_docx_path:{self.temp_pdf_path}")

        # 在预览区域显示 PDF
        pdf_url = QUrl.fromLocalFile(self.temp_pdf_path)
        print("PDF 文件 URL:", pdf_url.toString())
        self.preview.setUrl(pdf_url)

        # 检查 QWebEngineView 是否加载成功
        self.preview.loadFinished.connect(self.on_load_finished)
        QMessageBox.information(self, "成功", "预览已更新")

    def fill_word_template(self):
        """
        填充 Word 文档中的占位符并保存为新的 Word 文件
        :self.template_document_path: 模板 Word 文件路径
        """
        doc = Document(self.template_document_path)

        # 获取输出文件路径
        output_path, _ = QFileDialog.getSaveFileName(self, "保存文档", "", "PDF 文件 (*.pdf);;Word 文件 (*.docx)") # 获取输出文件的名称
        print(f"output_path is :{output_path}")

        # 替换段落中的占位符
        for paragraph in doc.paragraphs:
            for field, input_box in self.inputs.items():
                placeholder = field
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, input_box.text())

        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, input_box in self.inputs.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, input_box.text())

        # 保存填充后的 Word 文件
        if output_path.endswith(".docx"):
            doc.save(output_path)
            QMessageBox.information(self, "成功", "文档生成成功。")

        else:
            # 先转换成word，再转换成pdf
            temp_docx_path = os.path.join(tempfile.gettempdir(), "temp_template.docx")
            print("临时 Word 文件路径:", temp_docx_path)
            doc.save(temp_docx_path)
            convert(temp_docx_path, output_path)

            # 删除临时文件
            try:
                os.remove(temp_docx_path)
            except Exception as e:
                QMessageBox.information(self, "失败", "删除临时文件失败。")
                return

            QMessageBox.information(self, "成功", "文档生成成功。")




if __name__ == "__main__":
    # 确保 QApplication 在任何 QWidget 之前被创建
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())