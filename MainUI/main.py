import sys
import os
import tempfile
import warnings
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QVBoxLayout, QScrollArea, QWidget, QLabel,
    QLineEdit, QPushButton, QHBoxLayout, QMenuBar, QFileDialog, QMessageBox, QStackedWidget, QInputDialog, QDialog,
    QAction, QFormLayout
)
from PyQt5.QtCore import Qt, QUrl  # 导入 QUrl
from PyQt5.QtGui import QFont, QDoubleValidator, QIcon
from PyQt5.QtWebEngineWidgets import QWebEngineView, QWebEngineSettings  # 用于显示 PDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx2pdf import convert  # 用于将 Word 转换为 PDF
import re
from PyQt5.QtCore import QLoggingCategory
import pandas as pd  # 用于处理 Excel 文件
from CustomTools import CustomIntInputDialog, CustomTextInputDialog, number_to_rmb_upper, format_number_with_commas
import os
import win32com.client
import logging
import openpyxl 
from contextlib import redirect_stdout, redirect_stderr


logging.basicConfig(format='%(asctime)s - %(filename)s[line:%(lineno)d] - %(levelname)s: %(message)s',
                    level=logging.DEBUG,
                    filename='progress.log',
                    filemode='a')

QLoggingCategory.setFilterRules("js=false")

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        # 获取日志记录
        self.logger = logging.getLogger('progress')
        self.logger.info(f"程序的路径为：{os.getcwd()}")

        # 设置窗口标题和初始大小
        self.setWindowTitle("自动文档填充")
        self.setGeometry(100, 100, 1600, 900)

        # 设置icon图标
        self.setWindowIcon(QIcon("./icon/logo.png"))

        # 初始化 UI
        self.init_ui()

        # 初始状态隐藏主页面
        self.central_widget.hide()

    def init_ui(self):
        # 添加菜单栏
        menu_bar = QMenuBar(self)
        # 添加文件项
        file_menu = menu_bar.addMenu("文件")
        load_template_action = file_menu.addAction("合同模板")
        load_template_action.triggered.connect(self.load_template)

        # 添加 Excel 表单填写功能
        excel_form_action = file_menu.addAction("Excel 表单填写")
        excel_form_action.triggered.connect(self.show_excel_form_page)

        # 创建帮助 QAction
        help_action = QAction("帮助", self)
        help_action.triggered.connect(self.show_help_page)  # 绑定槽函数
        menu_bar.addAction(help_action)  # 将帮助 QAction 添加到菜单栏

        # 设置menuBar
        self.setMenuBar(menu_bar)

        # 创建主布局
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        main_layout = QHBoxLayout(self.central_widget)

        # 使用 QStackedWidget 来管理不同的页面
        self.stacked_widget = QStackedWidget()
        main_layout.addWidget(self.stacked_widget)

        # 创建主页面
        self.main_page = QWidget()
        self.main_layout = QHBoxLayout(self.main_page)

        # 左侧固定宽度区域（包含表单和提交按钮）
        self.left_widget = QWidget()
        self.left_widget.setFixedWidth(500)  # 固定宽度为 500
        self.left_layout = QVBoxLayout(self.left_widget)

        # 创建滚动视图
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.form_container = QWidget()
        self.form_layout = QVBoxLayout(self.form_container)

        # 添加表单内容到滚动视图
        self.inputs = {}
        self.scroll_area.setWidget(self.form_container)

        # 将滚动视图添加到左侧布局
        self.left_layout.addWidget(self.scroll_area)

        # 创建按钮布局（水平布局）
        button_layout = QHBoxLayout()

        # 添加“更新预览”按钮
        self.update_preview_button = QPushButton("更新预览")
        self.update_preview_button.setFont(QFont("Segoe UI", 10))
        self.update_preview_button.setStyleSheet(
            "QPushButton {"
            "padding: 10px;"
            "background-color: #28a745;"
            "color: white;"
            "border: none;"
            "border-radius: 4px;"
            "}"
            "QPushButton:hover {"
            "background-color: #218838;"
            "}"
        )
        self.update_preview_button.clicked.connect(self.update_preview)
        button_layout.addWidget(self.update_preview_button)

        # 添加"生成文档"按钮
        self.generate_button = QPushButton("生成文档")
        self.generate_button.setFont(QFont("Segoe UI", 10))
        self.generate_button.setStyleSheet(
            "QPushButton {"
            "padding: 10px;"
            "background-color: #007BFF;"
            "color: white;"
            "border: none;"
            "border-radius: 4px;"
            "}"
            "QPushButton:hover {"
            "background-color: #0056b3;"
            "}"
        )
        self.generate_button.clicked.connect(self.fill_word_template)
        button_layout.addWidget(self.generate_button)

        # 将按钮布局添加到左侧布局
        self.left_layout.addLayout(button_layout)

        # 右侧预览区域（PDF 预览）
        self.preview = QWebEngineView()  # 使用 QWebEngineView 显示 PDF
        self.preview.setStyleSheet(
            "QWebEngineView {"
            "padding: 10px;"
            "border: 1px solid #ccc;"
            "border-radius: 4px;"
            "background-color: #f5f5f5;"
            "}"
        )

        # 启用 PDF 插件
        self.preview.settings().setAttribute(QWebEngineSettings.PluginsEnabled, True)

        # 将左侧和右侧添加到主页面布局
        self.main_layout.addWidget(self.left_widget)
        self.main_layout.addWidget(self.preview)

        # 创建帮助页面
        self.help_page = QWidget()
        help_layout = QVBoxLayout(self.help_page)
        help_label = QLabel("这是帮助页面。\n\n你可以在这里提供一些使用说明或帮助信息。")
        help_label.setFont(QFont("Segoe UI", 12))
        help_label.setAlignment(Qt.AlignCenter)
        help_layout.addWidget(help_label)

        # 创建 Excel 表单填写页面
        self.excel_form_page = QWidget()
        excel_form_layout = QVBoxLayout(self.excel_form_page)

        # 添加表单布局
        self.excel_form = QFormLayout()
        self.excel_form.setSpacing(20)
        self.excel_inputs = []
        self.excel_field_names = ["产品单价", "产品数量", "税", "运费", "税", "税", "利润"]

        # 添加输入字段
        for i in range(len(self.excel_field_names)):
            label = QLabel(f"{self.excel_field_names[i]}")
            label.setFont(QFont("Segoe UI", 10))
            input_box = QLineEdit()
 
            # 使用 QDoubleValidator 允许输入小数
            validator = QDoubleValidator()  # 创建 QDoubleValidator
            validator.setDecimals(6)  # 设置允许的小数位数（例如 6 位小数）
            input_box.setValidator(validator)  # 应用验证器
            input_box.setFont(QFont("Segoe UI", 10))
            input_box.setStyleSheet(
                "QLineEdit {"
                "padding: 8px;"
                "border: 1px solid #ccc;"
                "border-radius: 4px;"
                "background-color: #f9f9f9;"
                "}"
                "QLineEdit:focus {"
                "border: 1px solid #007BFF;"
                "background-color: #fff;"
                "}"
            )
            input_box.textChanged.connect(self.on_text_changed) # 设置处理函数
            self.excel_inputs.append((label, input_box))
            self.excel_form.addRow(label, input_box)

        self.result_label = QLabel("合计: ")
        self.result_label.setFont(QFont("Segoe UI", 10))
        self.result_box = QLabel("")
        self.result_box.setFont(QFont("Segoe UI", 10))
        self.excel_form.addRow(self.result_label, self.result_box)

        excel_form_layout.addLayout(self.excel_form)

         # 添加按钮布局（水平布局）
        button_layout = QHBoxLayout()

        # 添加“修改字段名称”按钮
        self.modify_fields_button = QPushButton("修改字段名称")
        self.modify_fields_button.setFont(QFont("Segoe UI", 10))
        self.modify_fields_button.setStyleSheet(
            "QPushButton {"
            "padding: 10px;"
            "background-color: #6c757d;"
            "color: white;"
            "border: none;"
            "border-radius: 4px;"
            "}"
            "QPushButton:hover {"
            "background-color: #5a6268;"
            "}"
        )
        self.modify_fields_button.clicked.connect(self.modify_field_names)
        button_layout.addWidget(self.modify_fields_button)

        # 添加“生成 Excel”按钮
        self.generate_excel_button = QPushButton("生成 Excel")
        self.generate_excel_button.setFont(QFont("Segoe UI", 10))
        self.generate_excel_button.setStyleSheet(
            "QPushButton {"
            "padding: 10px;"
            "background-color: #007BFF;"
            "color: white;"
            "border: none;"
            "border-radius: 4px;"
            "}"
            "QPushButton:hover {"
            "background-color: #0056b3;"
            "}"
        )
        self.generate_excel_button.clicked.connect(self.generate_excel)
        button_layout.addWidget(self.generate_excel_button)

        # 将按钮布局添加到页面底部
        excel_form_layout.addLayout(button_layout)

        # 将主页面、帮助页面和 Excel 表单页面添加到 QStackedWidget
        self.stacked_widget.addWidget(self.main_page)
        self.stacked_widget.addWidget(self.help_page)
        self.stacked_widget.addWidget(self.excel_form_page)

        # 初始显示主页面
        self.stacked_widget.setCurrentWidget(self.main_page)

    def on_text_changed(self):
        """当input_box发生变化时，实时的更新内容"""
        numbers = []
        for field, input_box in self.excel_inputs:
            if not input_box.text():
                return
            numbers.append(float(input_box.text()))

        # 计算最终报价
        total = ((numbers[0]*numbers[1])*numbers[2] + numbers[3])*numbers[4]*numbers[5]*numbers[6]
        self.result_box.setText(str(total))

    def modify_field_names(self):
        """修改字段名称"""
        dialog = CustomIntInputDialog("修改字段名称", "请输入字段编号 (1-7):", 1, 7, self)
        if dialog.exec_() == QDialog.Accepted:
            field_number = dialog.get_value()
            text_dialog = CustomTextInputDialog("修改字段名称", f"请输入字段 {field_number} 的新名称:", self)
            if text_dialog.exec_() == QDialog.Accepted:
                new_name = text_dialog.get_text()
                if new_name:
                    self.excel_field_names[field_number - 1] = new_name
                    self.excel_inputs[field_number - 1][0].setText(f"{new_name}")

    def show_excel_form_page(self):
        """显示 Excel 表单填写页面"""
        self.stacked_widget.setCurrentWidget(self.excel_form_page)
        self.central_widget.show()

    def generate_excel(self):
        """生成 Excel 文件"""
        # 获取用户输入的值
        output_excel_Values = [input_box.text() for _, input_box in self.excel_inputs]
        if any(value == "" for value in output_excel_Values):
            QMessageBox.warning(self, "错误", "所有字段都必须填写数字。")
            return

        # 将值转换为浮点数
        try:
            output_excel_Values = [float(value) for value in output_excel_Values]
        except ValueError:
            QMessageBox.warning(self, "错误", "请输入有效的数字。")
            return

        # 禁用整个应用
        QApplication.setOverrideCursor(Qt.WaitCursor)

        output_excel_fields = self.excel_field_names[:]
        output_excel_fields.append("最终报价")
        print(f"output_excel_fields len : {len(output_excel_fields)}")
        # 计算最终的报价
        total = ((output_excel_Values[0]*output_excel_Values[1])*output_excel_Values[2] + output_excel_Values[3])*output_excel_Values[4]*output_excel_Values[5]*output_excel_Values[6]

        output_excel_Values.append(total)
        print(f"output_excel_values len : {len(output_excel_Values)}")

        # 创建 DataFrame
        data = {
            "项目": output_excel_fields,
            "数值": output_excel_Values
        }
        df = pd.DataFrame(data)

        # 获取保存路径
        output_path, _ = QFileDialog.getSaveFileName(self, "保存 Excel 文件", "", "Excel 文件 (*.xlsx)")
        if output_path:
            try:
                df.to_excel(output_path, index=False)
                
            except Exception as e:
                # 启用整个应用
                QApplication.restoreOverrideCursor()
                QMessageBox.critical(self, "错误", f"保存文件时出错: {str(e)}")
                return

        else:
            # 启用整个应用
            QApplication.restoreOverrideCursor()
            QMessageBox.information(self, "提醒", "请输入有效文件存储路径。")
            return

        # 启用整个应用
        QApplication.restoreOverrideCursor()
        QMessageBox.information(self, "成功", "Excel 文件生成成功。")
        

    def show_help_page(self):
        """显示帮助页面"""
        self.stacked_widget.setCurrentWidget(self.help_page)
        self.central_widget.show()

    def load_template(self):
        """加载 Word 模板文件"""
        file_path, _ = QFileDialog.getOpenFileName(self, "打开 Word 模板", "", "Word 文件 (*.docx)")

        self.logger.info('打开word模板文件路径：' + file_path)
        if file_path:
            # 加载模板
            try:
                self.template_document = Document(file_path)
                self.logger.info("模板文件加载成功")
                self.template_document_path = file_path  # 记录模板文件的地址

            except Exception as e:
                self.logger.error('加载模板文件失败，模板文件路径为：' + file_path)
                QMessageBox.information(self, "加载模板失败，请检查文件类型", str(e))
                return

            # 禁用整个应用
            QApplication.setOverrideCursor(Qt.WaitCursor)

            # 动态生成表单字段
            self.generate_form_fields()

            # 生成 PDF 预览
            self.generate_preview_pdf()

            # 切换QStackWeigt的id
            self.stacked_widget.setCurrentWidget(self.main_page)

            # 显示主页面
            self.central_widget.show()

            # 启用整个应用
            QApplication.restoreOverrideCursor()

    def generate_form_fields(self):
        """从模板中读取占位符并动态生成表单字段"""
        # 清空现有表单
        for i in reversed(range(self.form_layout.count())):
            self.form_layout.itemAt(i).widget().setParent(None)
        self.inputs.clear()

        # 读取模板中的占位符
        pattern = re.compile(r"\{\{.*?\}\}")  # 匹配占位符的正则表达式
        placeholders = list()
        for paragraph in self.template_document.paragraphs:
            matches = pattern.findall(paragraph.text)
            for match in matches:
                if match not in placeholders:
                    placeholders.append(match)

        # 从表格中提取占位符
        for table in self.template_document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = pattern.findall(paragraph.text)
                        for match in matches:
                            if match not in placeholders:  # 避免重复
                                placeholders.append(match)

        # 动态生成表单字段
        font_label = QFont("Segoe UI", 10)
        font_input = QFont("Segoe UI", 10)

        for placeholder in placeholders:
            label = QLabel(placeholder[2:-2])
            label.setFont(font_label)
            input_box = QLineEdit()
            input_box.setFont(font_input)
            input_box.setStyleSheet(
                "QLineEdit {"
                "padding: 8px;"
                "border: 1px solid #ccc;"
                "border-radius: 4px;"
                "background-color: #f9f9f9;"
                "}"
                "QLineEdit:focus {"
                "border: 1px solid #007BFF;"
                "background-color: #fff;"
                "}"
            )
            self.inputs[placeholder] = input_box
            self.form_layout.addWidget(label)
            self.form_layout.addWidget(input_box)

    def generate_preview_pdf(self):
        """生成 PDF 并在预览区域显示"""
        if not hasattr(self, 'template_document'):
            print("错误：未加载模板文件。")
            return

        # 创建临时 Word 文件
        self.temp_pdf_path = os.path.join(tempfile.gettempdir(), "tempPreview.pdf")
        self.logger.info("临时 PDF 文件路径: " + (self.temp_pdf_path))
        self.logger.info("模板 docx 文件路径: " + (self.template_document_path))

        try:
            # 将 stdout 和 stderr 重定向到文件
            with open("conversion.log", "w") as log_file:
                with redirect_stdout(log_file), redirect_stderr(log_file):
                    convert(self.template_document_path, self. temp_pdf_path)
            self.logger.info(f"模板 docx 文件{self.template_document_path} 转换为 PDF 成功")
        except Exception as e:
            print("Word 转换为 PDF 失败:", str(e))
            QMessageBox.information(self, "生成预览PDF失败", str(e))
            self.logger.error(f"模板 docx 文件{self.template_document_path} 转换为 PDF 失败:" + str(e))
            return

        # 检查 PDF 文件是否存在
        if not os.path.exists(self.temp_pdf_path):
            self.logger.error("错误：PDF 文件未生成。")
            return

        # 在预览区域显示 PDF
        pdf_url = QUrl.fromLocalFile(self.temp_pdf_path)
        self.preview.setUrl(pdf_url)

        # 检查 QWebEngineView 是否加载成功
        self.preview.loadFinished.connect(self.on_load_finished)

    def on_load_finished(self, success):
        """PDF 加载完成时的回调函数"""
        if success:
            print("PDF 文件加载回调成功。")
        else:
            print("PDF 文件加载回调失败。")

    def update_preview(self):
        """更新预览"""
        # 让程序进入等待状态
        # 禁用整个应用
        QApplication.setOverrideCursor(Qt.WaitCursor)

        # 第一步，打开document文件
        # 第二步，替换占位符
        # 第三步，生成临时文件word
        # 第四步，展示临时文件PDF
        doc = Document(self.template_document_path)

        # 替换段落中的占位符
        for paragraph in doc.paragraphs:
            for field, input_box in self.inputs.items():
                placeholder = field
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, format_number_with_commas(input_box.text()))
                    
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, input_box in self.inputs.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, format_number_with_commas(input_box.text()))
                            if "买方" in cell.text:
                                for paragraph in cell.paragraphs:
                                    # 文本加粗
                                    for run in paragraph.runs:
                                        run.bold = True
                                    # 左对齐
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                            else:
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    if "$合计$" in cell.text:
                        for placeholder, input_box in self.inputs.items():
                            if placeholder == "{{合计}}":
                                cell.text = cell.text.replace("$合计$", number_to_rmb_upper(input_box.text()))
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


        # 先转换成临时word
        temp_docx_path = os.path.join(tempfile.gettempdir(), "temp_template.docx")
        self.logger.info(f"更新预览生成的temp_docx_path:{temp_docx_path}")
        doc.save(temp_docx_path)

        # 转换成临时PDF
        # 将 stdout 和 stderr 重定向到文件
        try:
            with open("conversion.log", "w") as log_file:
                with redirect_stdout(log_file), redirect_stderr(log_file):
                    convert(temp_docx_path, self.temp_pdf_path)
            self.logger.info(f"更新预览生成的temp_docx_path:{self.temp_pdf_path}")

        except Exception as e:
            self.logger.error(f"更新预览时，word转换失败: {e}")
            return

        # 在预览区域显示 PDF
        pdf_url = QUrl.fromLocalFile(self.temp_pdf_path)
        print("PDF 文件 URL:", pdf_url.toString())
        self.preview.setUrl(pdf_url)

        # 启用整个应用
        QApplication.restoreOverrideCursor()

        # 检查 QWebEngineView 是否加载成功
        self.preview.loadFinished.connect(self.on_load_finished)
        QMessageBox.information(self, "成功", "预览已更新")
        self.logger.info("更新预览成功")

    def fill_word_template(self):
        """
        填充 Word 文档中的占位符并保存为新的 Word 文件
        :self.template_document_path: 模板 Word 文件路径
        """

        doc = Document(self.template_document_path)

        # 获取输出文件路径
        output_path, _ = QFileDialog.getSaveFileName(self, "保存文档", "", "PDF 文件 (*.pdf);;Word 文件 (*.docx)")  # 获取输出文件的名称

        # 禁用整个应用
        QApplication.setOverrideCursor(Qt.WaitCursor)

        # 替换段落中的占位符
        for paragraph in doc.paragraphs:
            for field, input_box in self.inputs.items():
                placeholder = field
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, input_box.text())

        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, input_box in self.inputs.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, format_number_with_commas(input_box.text()))
                            if "买方" in cell.text:
                                for paragraph in cell.paragraphs:
                                    # 文本加粗
                                    for run in paragraph.runs:
                                        run.bold = True
                                    # 左对齐
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            else:
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    if "$合计$" in cell.text:
                        for placeholder, input_box in self.inputs.items():
                            if placeholder == "{{合计}}":
                                cell.text = cell.text.replace("$合计$", number_to_rmb_upper(input_box.text()))
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 保存填充后的 Word 文件
        if output_path.endswith(".docx"):
            doc.save(output_path)
            self.logger.info("docx文档保存成功")
            
        else:
            # 先转换成word，再转换成pdf
            temp_docx_path = os.path.join(tempfile.gettempdir(), "temp_template.docx")
            self.logger.info("保存文件时，临时 Word 文件路径: ", temp_docx_path)
            try:
                doc.save(temp_docx_path)
                # 将 stdout 和 stderr 重定向到文件
                with open("conversion.log", "w") as log_file:
                    with redirect_stdout(log_file), redirect_stderr(log_file):
                        convert(temp_docx_path, output_path)
                self.logger.info("pdf文档保存成功")
            except Exception as e:
                # 启用整个应用
                QApplication.restoreOverrideCursor()

                QMessageBox.information(self, "文档保存失败", str(e))
                self.logger.error("pdf文档保存失败")
                return
            # 删除临时文件
            try:
                os.remove(temp_docx_path)
            except Exception as e:
                # 启用整个应用
                QApplication.restoreOverrideCursor()

                QMessageBox.information(self, "失败", "删除临时文件失败。")
                self.logger.error("删除临时文件失败")
                return
            
        # 启用整个应用
        QApplication.restoreOverrideCursor()
        QMessageBox.information(self, "成功", "文档生成成功。")

# import ctypes
def setup_environment():
    """动态设置环境变量"""
    word_path = r"C:\Program Files\Microsoft Office\root\Office16"
    if os.path.exists(os.path.join(word_path, "WINWORD.EXE")):
        os.environ["PATH"] += os.pathsep + word_path
    else:
        print("未找到 Microsoft Word 的路径，转换可能失败")

if __name__ == "__main__":
    # 确保 QApplication 在任何 QWidget 之前被创建
    app = QApplication(sys.argv)

    #动态设置环境变量
    setup_environment()

    window = MainWindow()
    window.show()
    sys.exit(app.exec_())

