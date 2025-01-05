from docx import Document
from docx2pdf import convert

def fill_word_template(template_path, output_path, placeholders):
    """
    填充 Word 文档中的占位符并保存为新的 Word 文件
    :param template_path: 模板 Word 文件路径
    :param output_path: 填充后的 Word 文件路径
    :param placeholders: 占位符字典（如 {"{{name}}": "张三"}）
    """
    # 打开模板
    doc = Document(template_path)
    
    # 遍历每一段落
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    
    # 遍历表格（如果 Word 中有表格也需要替换）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placeholders.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)
    
    # 保存填充后的 Word 文件
    doc.save(output_path)

def convert_word_to_pdf(word_path, pdf_path):
    """
    将 Word 文件转换为 PDF 文件
    :param word_path: Word 文件路径
    :param pdf_path: 输出的 PDF 文件路径
    """
    # 调用 docx2pdf 转换
    convert(word_path, pdf_path)

# 示例用法
if __name__ == "__main__":
    # 模板路径
    template_path = r"C:\Users\11482\Desktop\test.docx"  # 模板文件
    output_word_path = "filled_output.docx"  # 填充后的 Word 文件
    output_pdf_path = "filled_output.pdf"  # 最终生成的 PDF 文件

    # 占位符字典
    placeholders = {
        "{{name}}": "张三",
    }

    # 填充模板
    fill_word_template(template_path, output_word_path, placeholders)

    # 转换为 PDF
    convert_word_to_pdf(output_word_path, output_pdf_path)

    print("PDF 文件生成成功！")
