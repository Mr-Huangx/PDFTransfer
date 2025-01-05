from docx import Document
from docx2pdf import convert
import re

def extract_placeholders_from_word(template_path):
    """
    提取 Word 文档中的所有占位符
    :param template_path: 模板 Word 文件路径
    :return: 占位符列表
    """
    doc = Document(template_path)
    placeholders = set()  # 使用集合去重
    pattern = re.compile(r"\{\{.*?\}\}")  # 匹配占位符的正则表达式，如 {{name}}

    # 提取段落中的占位符
    for paragraph in doc.paragraphs:
        matches = pattern.findall(paragraph.text)
        placeholders.update(matches)

    # 提取表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = pattern.findall(cell.text)
                placeholders.update(matches)

    return list(placeholders)

def fill_word_template(template_path, output_path, placeholders):
    """
    填充 Word 文档中的占位符并保存为新的 Word 文件
    :param template_path: 模板 Word 文件路径
    :param output_path: 填充后的 Word 文件路径
    :param placeholders: 占位符字典，如 {"{{name}}": "张三"}
    """
    doc = Document(template_path)

    # 替换段落中的占位符
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placeholders.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

    # 保存填充后的 Word 文件
    doc.save(output_path)

def convert_word_to_pdf(word_path, pdf_path):
    """
    将 Word 文件转换为 PDF 文件
    :param word_path: Word 文件路径
    :param pdf_path: 输出的 PDF 文件路径
    """
    # 调用 docx2pdf 转换
    convert(word_path, pdf_path)

# 示例用法
if __name__ == "__main__":
    # 模板路径
    template_path = r"C:\Users\11482\Desktop\test.docx"   # 模板文件
    output_word_path = "filled_output.docx"  # 填充后的 Word 文件
    output_pdf_path = "filled_output.pdf"  # 最终生成的 PDF 文件

    # 提取占位符
    placeholders = extract_placeholders_from_word(template_path)
    print("提取的占位符：", placeholders)

    # 示例占位符替换字典
    placeholder_values = {
        "{{name}}": "张三",
        "{{date}}": "2025-01-05",
        "{{address}}": "北京市海淀区"
    }

    # 填充模板
    fill_word_template(template_path, output_word_path, placeholder_values)

    # 转换为 PDF
    convert_word_to_pdf(output_word_path, output_pdf_path)

    print("PDF 文件生成成功！")
